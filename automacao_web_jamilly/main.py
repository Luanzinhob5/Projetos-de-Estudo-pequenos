from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as ec
from time import sleep
import os
from dotenv import load_dotenv
from docx import Document
import customtkinter as ctk

#Verifica se o nome do advogado esta no word
def analizar_se_ja_possui(adv_nome, documento):
    texto_completo = ""
    for paragrafo in documento.paragraphs:
        texto_completo += paragrafo.text + "/n"

    if adv_nome in texto_completo:
        return True
    else:
        return False


load_dotenv()

driver = webdriver.Chrome()

wait = WebDriverWait(driver, 20)

email = os.environ["EMAIL"]
senha = os.environ["SENHA"]

try:
    documento = Document("dados_adv_programa.docx")
except:
    documento = Document()

#Entra no site
driver.get("https://dataged.eastus.cloudapp.azure.com")

#Adicionar email ao login
campo_email = wait.until(ec.presence_of_element_located((By.ID,"email")))
campo_email.click()
campo_email.send_keys(email)

#Adicionar senha ao login
campo_senha = wait.until(ec.presence_of_element_located((By.ID,"senha")))
campo_senha.click()
campo_senha.send_keys(senha)
sleep(20)

#Clicar no botao de Login
botao_logar = wait.until(ec.presence_of_element_located((By.CLASS_NAME,"buttonGoogle")))
botao_logar.click()

#Entrar na caixa de analise
caixa_analise = wait.until(ec.presence_of_element_located((By.XPATH,'//*[@id="side-menu"]/li[1]/ul/li[2]/a')))
caixa_analise.click()

#Verificar documentos com certificado digital
links = driver.find_elements(By.XPATH, '//*[@id="normal"]/form/table/tbody/tr/td[1]/a')
total_links = len(links)

#Pega o indice de cada link
for i in range(total_links):
    links = driver.find_elements(By.XPATH, '//*[@id="normal"]/form/table/tbody/tr/td[1]/a')
    link = links[i]

    if link.find_element(By.XPATH, './../../td[4]').text == "CERTIFICADO DIGITAL":
        janela_principal = driver.current_window_handle
        nome_adv = link.find_element(By.XPATH, "./../../td[3]").text

        if analizar_se_ja_possui(nome_adv,documento):
            pass

        else:
            link.click()

            #Verifica se cada documento possui OAB no nome
            entrar_aba_documentos = wait.until(ec.presence_of_element_located((By.XPATH, '//*[@id="page-wrapper"]/div/div[2]/div/ul/li[3]/a')))
            entrar_aba_documentos.click()

            aba_documentos = driver.find_elements(By.XPATH, '//*[@id="listaDeDocs"]/table/tbody/tr')
            n = 0
            for documento_recebido in aba_documentos:
                if "OAB" in documento_recebido.find_element(By.XPATH, './td[2]').text.upper():
                    n += 1

            documentacao = f"documentos OAB: {n}."

            #Volta para a pagina inicial do Dataged
            voltar_inicio = wait.until(ec.presence_of_element_located((By.XPATH, '//*[@id="wrapper"]/nav/div[2]/a/img'))).click()

            #Clica no Ficha do Advogado
            ficha_advogado = wait.until(ec.presence_of_element_located((By.XPATH, '//*[@id="page-wrapper"]/div[4]/div/div/div[6]/form/button'))).click()
            janelas_abertas = driver.window_handles

            for janela in janelas_abertas:
                if janela not in janela_principal:
                    driver.switch_to.window(janela)
                    consulta_nome = wait.until(ec.presence_of_element_located((By.XPATH, '//*[@id="nome"]')))
                    consulta_nome.click()
                    consulta_nome.send_keys(nome_adv)

                    pesquisar = wait.until(ec.presence_of_element_located((By.XPATH, '//*[@id="Install2"]')))
                    pesquisar.click()
                    #Pegar os dados de email, cpf, data de nascimento, telefone, subsecao, oab do advogado.
                    try:
                        email = driver.find_element(By.XPATH, '//*[@id="frm"]/table[2]/tbody/tr[2]/td[8]').text
                        cpf = driver.find_element(By.XPATH, '//*[@id="frm"]/table[2]/tbody/tr[2]/td[5]').text
                        data_de_nascimento = driver.find_element(By.XPATH, '//*[@id="frm"]/table[2]/tbody/tr[2]/td[12]').text
                        telefone = driver.find_element(By.XPATH, '//*[@id="frm"]/table[2]/tbody/tr[2]/td[11]').text
                        subsecao = driver.find_element(By.XPATH, '//*[@id="frm"]/table[2]/tbody/tr[2]/td[14]').text
                        oab = driver.find_element(By.XPATH, '//*[@id="frm"]/table[2]/tbody/tr[2]/td[2]').text
                        texto = f"{documentacao}\n\nNome: {nome_adv}\nEmail: {email}\nCPF: {cpf}\nData de Nascimento: {data_de_nascimento}\nTelefone: {telefone}\nSubsecao: {subsecao}\nOAB: {oab}\n\n\n"
                        documento.add_paragraph(texto)
                        documento.save("dados_adv_programa.docx")
                    except:
                        texto = f"{documentacao}\n\nNome: {nome_adv}\nERRO AO PEGAR email/cpf/ddn/tel/sub/oab\n\nFAZER MANUALMENTE\n\n\n"
                        documento.add_paragraph(texto)
                        documento.save("dados_adv_programa.docx")

                    driver.close()
            
        driver.switch_to.window(janela_principal)
        caixa_analise2 = wait.until(ec.presence_of_element_located((By.XPATH, '//*[@id="side-menu"]/li[1]/ul/li[2]/a')))
        caixa_analise2.click()

    


															


